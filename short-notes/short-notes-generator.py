"""
NIE Sri Lanka — Short Notes Generator
======================================
Reads downloaded PDF textbooks/guides and generates structured short notes
per lesson as a Word (.docx) document using DeepSeek AI.

Prerequisites:
    pip install openai pdfplumber python-docx

Usage:
    # Single PDF
    python short-notes-generator.py --pdf "path/to/Science.pdf" --grade 10 --api-key "sk-..."

    # All PDFs in a folder (e.g. one subject with multiple chapter PDFs)
    python short-notes-generator.py --folder "../../output/textbooks/grade 10 -en/Science Part I" --grade 10 --api-key "sk-..."

    # Use DEEPSEEK_API_KEY environment variable instead
    python short-notes-generator.py --pdf "..." --grade 10

    # Custom output folder
    python short-notes-generator.py --pdf "..." --grade 10 --output "output/"
"""

import os, sys, re, json, argparse, time
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

try:
    from openai import OpenAI
except ImportError:
    print("❌ Missing: pip install openai"); sys.exit(1)

try:
    import pdfplumber
except ImportError:
    print("❌ Missing: pip install pdfplumber"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ Missing: pip install python-docx"); sys.exit(1)

# ── PDF extraction ─────────────────────────────────────────────────────────────

def extract_pdf_text(pdf_path, max_pages=None):
    """Extract text from a PDF file using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            pages = pdf.pages[:max_pages] if max_pages else pdf.pages
            for i, page in enumerate(pages, 1):
                t = page.extract_text()
                if t and t.strip():
                    text_parts.append(f"[Page {i}]\n{t.strip()}")
    except Exception as e:
        print(f"   ⚠️  Could not extract text from {pdf_path}: {e}")
        return ""
    return "\n\n".join(text_parts)

def collect_pdfs(source):
    """Return list of PDF paths from a file or folder."""
    p = Path(source)
    if p.is_file() and p.suffix.lower() == '.pdf':
        return [p]
    if p.is_dir():
        pdfs = sorted(p.rglob('*.pdf'))
        if not pdfs:
            print(f"   ❌ No PDFs found in {source}")
            sys.exit(1)
        return pdfs
    print(f"   ❌ Not a valid file or folder: {source}")
    sys.exit(1)

# ── DeepSeek API ──────────────────────────────────────────────────────────────

NOTES_TOOL = {
    "type": "function",
    "function": {
        "name": "create_short_notes",
        "description": "Create structured short notes from lesson/chapter content.",
        "parameters": {
            "type": "object",
            "required": ["subject_title", "lessons"],
            "properties": {
                "subject_title": {
                    "type": "string",
                    "description": "Full subject name as identified from the content"
                },
                "lessons": {
                    "type": "array",
                    "description": "One entry per lesson/chapter identified in the content",
                    "items": {
                        "type": "object",
                        "required": ["lesson_number", "lesson_title", "introduction",
                                     "core_concepts", "key_processes", "exam_tips",
                                     "important_points"],
                        "properties": {
                            "lesson_number": {"type": "integer"},
                            "lesson_title":  {"type": "string"},
                            "introduction": {
                                "type": "object",
                                "required": ["text", "bullet_points"],
                                "properties": {
                                    "text":          {"type": "string",
                                                      "description": "1-3 sentence intro paragraph"},
                                    "bullet_points": {"type": "array", "items": {"type": "string"},
                                                      "description": "3-6 key introductory facts"}
                                }
                            },
                            "core_concepts": {
                                "type": "array",
                                "description": "2-5 major concept groups (A, B, C…)",
                                "items": {
                                    "type": "object",
                                    "required": ["label", "title", "text", "bullet_points"],
                                    "properties": {
                                        "label":         {"type": "string",
                                                          "description": "e.g. A, B, C"},
                                        "title":         {"type": "string"},
                                        "text":          {"type": "string",
                                                          "description": "Short explanatory paragraph"},
                                        "bullet_points": {"type": "array", "items": {"type": "string"}}
                                    }
                                }
                            },
                            "key_processes": {
                                "type": "array",
                                "description": "Definitions/processes students must memorize",
                                "items": {
                                    "type": "object",
                                    "required": ["name", "explanation"],
                                    "properties": {
                                        "name":        {"type": "string"},
                                        "explanation": {"type": "string"}
                                    }
                                }
                            },
                            "exam_tips": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Practical tips for answering exam questions"
                            },
                            "important_points": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Must-know facts, safety notes, or common mistakes"
                            }
                        }
                    }
                }
            }
        }
    }
}

SYSTEM_PROMPT = """\
You are an expert Sri Lankan education content creator specialising in NIE curriculum.
Given raw textbook/guide content, you:
1. Identify EVERY lesson or chapter in the content.
2. Create comprehensive yet concise short notes for EACH lesson.
3. Write for students preparing for Grade 6-13 national exams.
4. Use simple, clear English.
5. Include all important facts, definitions, examples, and exam hints.
6. Always call the create_short_notes function with your result - do NOT write prose.
"""

def generate_notes_via_deepseek(text, grade, subject_hint, api_key):
    """Call DeepSeek API and return the structured notes JSON."""
    client = OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com"
    )

    user_msg = f"""Grade: {grade}
Subject hint: {subject_hint or 'Unknown - identify from content'}

Below is the raw extracted text from the PDF(s). Identify all lessons/chapters and generate
complete short notes for each one using the create_short_notes function.

--- BEGIN CONTENT ---
{text[:120000]}
--- END CONTENT ---"""

    print("   🤖 Sending to DeepSeek... ", end='', flush=True)
    t0 = time.time()

    response = client.chat.completions.create(
        model="deepseek-chat",
        max_tokens=8192,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": user_msg}
        ],
        tools=[NOTES_TOOL],
        tool_choice={"type": "function", "function": {"name": "create_short_notes"}}
    )

    elapsed = time.time() - t0
    print(f"done ({elapsed:.1f}s)")

    # Extract function call result
    msg = response.choices[0].message
    if msg.tool_calls:
        return json.loads(msg.tool_calls[0].function.arguments)

    raise ValueError("DeepSeek did not return the expected function call.")

# ── Word document builder ──────────────────────────────────────────────────────

def _set_color(run, hex_color):
    """Set run font colour using hex string e.g. '1A237E'."""
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )

def _heading(doc, text, level, color_hex=None, bold=True):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.bold = bold
        if color_hex:
            _set_color(run, color_hex)
    return p

def _body(doc, text, bold_prefix=None):
    """Add a body paragraph; optionally bold the text before the first colon."""
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix and ':' in text:
        idx = text.index(':') + 1
        run_bold = p.add_run(text[:idx])
        run_bold.bold = True
        run_rest = p.add_run(text[idx:])
    else:
        p.add_run(text)

def _bullet(doc, text, level=0):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(1)
    # Bold any text before a colon at the start
    if ':' in text[:50]:
        idx = text.index(':') + 1
        run_b = p.add_run(text[:idx])
        run_b.bold = True
        p.add_run(text[idx:])
    else:
        p.add_run(text)

def _divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BDBDBD')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def build_word_doc(notes_data, grade, output_path):
    """Build a formatted Word document from the structured notes data."""
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Default body font ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    subject_title = notes_data.get('subject_title', 'Subject')

    # ── Title page block ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"Grade {grade} — {subject_title}")
    run.bold = True
    run.font.size = Pt(20)
    _set_color(run, '1A237E')

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Short Notes  |  NIE Sri Lanka Curriculum")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph()  # spacer

    # ── Lessons ──
    lessons = notes_data.get('lessons', [])
    for lesson in lessons:
        ln  = lesson.get('lesson_number', '?')
        lt  = lesson.get('lesson_title', 'Lesson')

        # Lesson heading
        _heading(doc, f"Lesson {ln}: {lt}", level=1, color_hex='1A237E')

        # ── 1. Introduction ──
        _heading(doc, f"1. Introduction to {lt}", level=2, color_hex='283593')
        intro = lesson.get('introduction', {})
        if intro.get('text'):
            _body(doc, intro['text'])
        for bp in intro.get('bullet_points', []):
            _bullet(doc, bp)

        # ── 2. Core Concepts ──
        concepts = lesson.get('core_concepts', [])
        if concepts:
            _heading(doc, "2. Core Concepts", level=2, color_hex='283593')
            for concept in concepts:
                label = concept.get('label', '')
                title = concept.get('title', '')
                _heading(doc, f"{label}. {title}", level=3, color_hex='37474F')
                if concept.get('text'):
                    _body(doc, concept['text'])
                for bp in concept.get('bullet_points', []):
                    _bullet(doc, bp)

        # ── 3. Key Processes ──
        processes = lesson.get('key_processes', [])
        if processes:
            _heading(doc, "3. Key Processes to Memorize", level=2, color_hex='283593')
            for proc in processes:
                _body(doc, f"{proc.get('name','')}: {proc.get('explanation','')}", bold_prefix=True)

        # ── 4. Tips & Tricks ──
        tips = lesson.get('exam_tips', [])
        if tips:
            _heading(doc, "4. Tips & Tricks for Exams", level=2, color_hex='283593')
            for tip in tips:
                _bullet(doc, tip)

        # ── 5. Important Points ──
        points = lesson.get('important_points', [])
        if points:
            _heading(doc, "5. Important Points to Remember", level=2, color_hex='283593')
            for pt in points:
                _bullet(doc, pt)

        _divider(doc)
        doc.add_paragraph()

    doc.save(output_path)
    print(f"   📄 Saved: {output_path}")

# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="NIE Sri Lanka — Short Notes Generator (Claude AI + Word)")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--pdf',    '-f', help='Path to a single PDF file')
    group.add_argument('--folder', '-F', help='Folder containing PDF file(s)')
    parser.add_argument('--grade',   '-g', required=True, help='Grade number (e.g. 10)')
    parser.add_argument('--subject', '-s', default=None,
                        help='Subject name hint (optional, auto-detected if omitted)')
    parser.add_argument('--output',  '-o', default='output',
                        help='Output folder for .docx files (default: output/)')
    parser.add_argument('--api-key', '-k', default=None,
                        help='DeepSeek API key (or set DEEPSEEK_API_KEY env variable)')
    parser.add_argument('--max-pages', type=int, default=None,
                        help='Limit pages extracted per PDF (useful for very large files)')
    args = parser.parse_args()

    api_key = args.api_key or os.environ.get('DEEPSEEK_API_KEY')
    if not api_key:
        print("❌ No API key found. Use --api-key or set DEEPSEEK_API_KEY environment variable.")
        sys.exit(1)

    # Collect PDFs
    source = args.pdf or args.folder
    pdfs = collect_pdfs(source)

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("📝 NIE Short Notes Generator")
    print("=" * 60)
    print(f"   Grade:   {args.grade}")
    print(f"   PDFs:    {len(pdfs)} file(s)")
    print(f"   Output:  {output_dir.resolve()}")
    print("=" * 60)

    # Extract text from all PDFs
    print("\n📖 Extracting text from PDF(s)...")
    all_text_parts = []
    for pdf in pdfs:
        print(f"   • {pdf.name}")
        text = extract_pdf_text(pdf, max_pages=args.max_pages)
        if text:
            all_text_parts.append(f"=== {pdf.stem} ===\n{text}")
        else:
            print(f"   ⚠️  No text extracted (may be a scanned image PDF)")

    if not all_text_parts:
        print("❌ No text could be extracted from any PDF.")
        print("   Note: Scanned image PDFs require OCR and are not currently supported.")
        sys.exit(1)

    combined_text = "\n\n".join(all_text_parts)
    word_count = len(combined_text.split())
    print(f"   ✅ Extracted ~{word_count:,} words from {len(all_text_parts)} PDF(s)")

    # Derive subject hint from folder/file name if not provided
    subject_hint = args.subject
    if not subject_hint:
        subject_hint = Path(source).stem if args.pdf else Path(source).name

    # Call Claude API
    print(f"\n🤖 Generating short notes with DeepSeek AI...")
    try:
        notes_data = generate_notes_via_deepseek(
            combined_text, args.grade, subject_hint, api_key)
    except Exception as e:
        print(f"\n❌ DeepSeek API error: {e}")
        sys.exit(1)

    subject_title = notes_data.get('subject_title', subject_hint or 'Subject')
    n_lessons = len(notes_data.get('lessons', []))
    print(f"   ✅ Generated notes for {n_lessons} lesson(s) in '{subject_title}'")

    # Build Word document
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', subject_title)
    out_file = output_dir / f"Grade {args.grade} - {safe_name} - Short Notes.docx"

    print(f"\n📄 Building Word document...")
    build_word_doc(notes_data, args.grade, str(out_file))

    # Save raw JSON alongside (useful for debugging / re-formatting)
    json_file = output_dir / f"Grade {args.grade} - {safe_name} - Short Notes.json"
    with open(json_file, 'w', encoding='utf-8') as jf:
        json.dump(notes_data, jf, ensure_ascii=False, indent=2)
    print(f"   🗂️  JSON:  {json_file}")

    print(f"\n{'=' * 60}")
    print(f"✅ Done!  {n_lessons} lesson(s)  →  {out_file.name}")
    print(f"{'=' * 60}")

if __name__ == '__main__':
    main()
