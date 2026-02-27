"""
NIE Sri Lanka — Short Notes Generator (Expanded 6-Section Format)
==================================================================
Reads NIE teacher's guides (and optionally textbook PDFs) and generates
structured short notes per lesson as individual Word (.docx) files.

Lesson divisions are taken from the TEACHER'S GUIDE, which defines lesson
numbers, titles, objectives, and content scope. Textbook PDFs are optional
supplementary content used to enrich the notes.

Output structure:
    <output>/<Subject Name>/Lesson 01 - Title.docx
    <output>/<Subject Name>/Lesson 02 - Title.docx
    ...

Prerequisites:
    pip install openai pdfplumber python-docx

Usage:
    # Teacher's guide only (lesson structure + content from guide)
    python short-notes-generator.py --guide "output/teachers-guides/grade 10 -en/Mathematics/Mathematics.pdf" --grade 10

    # Teacher's guide folder
    python short-notes-generator.py --guide-folder "output/teachers-guides/grade 10 -en/Mathematics/" --grade 10

    # Teacher's guide + textbook for richer notes
    python short-notes-generator.py \
        --guide "output/teachers-guides/grade 10 -en/Science/Science.pdf" \
        --pdf   "output/textbooks/grade 10 -en/Science Part I/Science.pdf" \
        --grade 10

    # Teacher's guide folder + textbook folder
    python short-notes-generator.py \
        --guide-folder "output/teachers-guides/grade 10 -en/Mathematics/" \
        --folder       "output/textbooks/grade 10 -en/Mathematics/" \
        --grade 10

    # Set API key via environment variable
    export DEEPSEEK_API_KEY="sk-..."
    python short-notes-generator.py --guide "..." --grade 10
"""

import os, sys, re, json, argparse
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

try:
    from openai import OpenAI
except ImportError:
    print("❌ Missing: pip install openai"); sys.exit(1)

try:
    import pdfplumber
except ImportError:
    print("❌ Missing: pip install pdfplumber"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ Missing: pip install python-docx"); sys.exit(1)

# ── PDF extraction ──────────────────────────────────────────────────────────────

def extract_pdf_text(pdf_path, max_pages=None):
    text_parts = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            pages = pdf.pages[:max_pages] if max_pages else pdf.pages
            for i, page in enumerate(pages, 1):
                t = page.extract_text()
                if t and t.strip():
                    text_parts.append(f"[Page {i}]\n{t.strip()}")
    except Exception as e:
        print(f"   ⚠️  Could not extract text from {pdf_path}: {e}")
        return ""
    return "\n\n".join(text_parts)

def collect_pdfs(source):
    p = Path(source)
    if p.is_file() and p.suffix.lower() == '.pdf':
        return [p]
    if p.is_dir():
        pdfs = sorted(p.rglob('*.pdf'))
        if not pdfs:
            print(f"   ❌ No PDFs found in {source}"); sys.exit(1)
        return pdfs
    print(f"   ❌ Not a valid file or folder: {source}"); sys.exit(1)

# ── DeepSeek API tool schema ────────────────────────────────────────────────────

NOTES_TOOL = {
    "type": "function",
    "function": {
        "name": "create_short_notes",
        "description": (
            "Create detailed structured short notes for every lesson/chapter "
            "in the given content. English only. Medium detail (2-3 pages per lesson)."
        ),
        "parameters": {
            "type": "object",
            "required": ["subject_title", "lessons"],
            "properties": {
                "subject_title": {
                    "type": "string",
                    "description": "Full subject name identified from content"
                },
                "lessons": {
                    "type": "array",
                    "description": "One entry per lesson/chapter — cover ALL lessons found",
                    "items": {
                        "type": "object",
                        "required": [
                            "lesson_number", "lesson_title",
                            "core_concepts", "key_formulas",
                            "worked_examples", "diagram_descriptions",
                            "exam_tips", "important_points"
                        ],
                        "properties": {
                            "lesson_number": {"type": "integer"},
                            "lesson_title":  {"type": "string"},

                            # ── Section 1: Core Concepts ──────────────────────
                            "core_concepts": {
                                "type": "array",
                                "description": "2-5 major concept groups, each with full detail",
                                "items": {
                                    "type": "object",
                                    "required": [
                                        "label", "title", "definition",
                                        "explanation", "analogy",
                                        "diagram_note", "connections"
                                    ],
                                    "properties": {
                                        "label":       {"type": "string",
                                                        "description": "A, B, C..."},
                                        "title":       {"type": "string"},
                                        "definition":  {"type": "string",
                                                        "description": "Formal/textbook definition"},
                                        "explanation": {"type": "string",
                                                        "description": "Simple plain-language explanation (2-4 sentences)"},
                                        "analogy":     {"type": "string",
                                                        "description": "Real-world analogy to aid memory"},
                                        "diagram_note":{"type": "string",
                                                        "description": "Reference to any related diagram or visual"},
                                        "connections": {"type": "array",
                                                        "items": {"type": "string"},
                                                        "description": "Links to other concepts in this lesson or subject"}
                                    }
                                }
                            },

                            # ── Section 2: Key Formulas / Definitions / Rules ─
                            "key_formulas": {
                                "type": "object",
                                "description": "Split into Revision (previously learned) and New (introduced in this lesson)",
                                "required": ["revision", "new"],
                                "properties": {
                                    "revision": {
                                        "type": "array",
                                        "items": {
                                            "type": "object",
                                            "required": ["name", "formula_or_rule", "worked_example"],
                                            "properties": {
                                                "name":            {"type": "string"},
                                                "formula_or_rule": {"type": "string"},
                                                "worked_example":  {
                                                    "type": "string",
                                                    "description": "Full step-by-step substitution / application example"
                                                }
                                            }
                                        }
                                    },
                                    "new": {
                                        "type": "array",
                                        "items": {
                                            "type": "object",
                                            "required": ["name", "formula_or_rule", "worked_example"],
                                            "properties": {
                                                "name":            {"type": "string"},
                                                "formula_or_rule": {"type": "string"},
                                                "worked_example":  {"type": "string"}
                                            }
                                        }
                                    }
                                }
                            },

                            # ── Section 3: Worked Examples ────────────────────
                            "worked_examples": {
                                "type": "array",
                                "description": "2-4 fully solved problems with step-by-step solutions",
                                "items": {
                                    "type": "object",
                                    "required": [
                                        "problem_number", "problem",
                                        "steps", "tricky_parts", "common_mistakes"
                                    ],
                                    "properties": {
                                        "problem_number":  {"type": "integer"},
                                        "problem":         {"type": "string",
                                                            "description": "The full question/problem statement"},
                                        "steps":           {"type": "array",
                                                            "items": {"type": "string"},
                                                            "description": "Numbered steps showing complete solution"},
                                        "tricky_parts":    {"type": "array",
                                                            "items": {"type": "string"},
                                                            "description": "Highlighted parts students often get wrong"},
                                        "common_mistakes": {"type": "array",
                                                            "items": {"type": "string"}}
                                    }
                                }
                            },

                            # ── Section 4: Diagrams & Visual Descriptions ─────
                            "diagram_descriptions": {
                                "type": "array",
                                "description": "Describe every diagram/figure in the source material",
                                "items": {
                                    "type": "object",
                                    "required": [
                                        "diagram_number", "title",
                                        "what_it_shows", "labels",
                                        "relationships", "key_learning"
                                    ],
                                    "properties": {
                                        "diagram_number": {"type": "integer"},
                                        "title":          {"type": "string"},
                                        "what_it_shows":  {"type": "string",
                                                           "description": "What the diagram illustrates"},
                                        "labels":         {"type": "array",
                                                           "items": {"type": "string"},
                                                           "description": "List of all labels/parts shown"},
                                        "relationships":  {"type": "string",
                                                           "description": "How parts relate to each other"},
                                        "key_learning":   {"type": "string",
                                                           "description": "What students must understand from this diagram"}
                                    }
                                }
                            },

                            # ── Section 5: Tips & Tricks for Exams ───────────
                            "exam_tips": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "required": ["tip", "mark_info", "frequency"],
                                    "properties": {
                                        "tip":       {"type": "string",
                                                      "description": "Practical exam tip"},
                                        "mark_info": {"type": "string",
                                                      "description": "Mark allocation, e.g. '2 marks', 'essay 10 marks'"},
                                        "frequency": {"type": "string",
                                                      "description": "How often this appears, e.g. 'every year', 'occasionally'"}
                                    }
                                }
                            },

                            # ── Section 6: Important Points to Remember ───────
                            "important_points": {
                                "type": "object",
                                "required": ["must_know", "exam_checklist"],
                                "properties": {
                                    "must_know": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                        "description": "Key facts, definitions, and safety notes"
                                    },
                                    "exam_checklist": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                        "description": "Checklist items students must verify before answering"
                                    }
                                }
                            }
                        }
                    }
                }
            }
        }
    }
}

SYSTEM_PROMPT = """\
You are an expert Sri Lankan education content creator specialising in the NIE curriculum.

You will receive content from a TEACHER'S GUIDE and optionally a TEXTBOOK.
Your responsibilities:

1. Use the TEACHER'S GUIDE as the authoritative source for lesson structure.
   - Identify every lesson by its lesson number, title, and scope as defined in the guide.
   - The guide defines the official lesson divisions — do NOT merge or split lessons.
   - Use the guide's learning outcomes and content outline for each lesson.

2. Enrich notes using TEXTBOOK CONTENT where provided.
   - Cross-reference textbook examples, diagrams, and explanations for each lesson.
   - If no textbook is provided, rely entirely on the teacher's guide content.

3. Create detailed short notes for EVERY lesson — English only.
   - Target medium detail: approximately 2-3 pages per lesson.
   - Always call create_short_notes — never write plain prose.

For each lesson produce all 6 sections:
  Section 1 — Core Concepts: full definition, plain-language explanation, real-world analogy,
               diagram reference, and connections to other concepts.
  Section 2 — Key Formulas/Definitions/Rules: split into Revision (previously learned) and
               New (introduced this lesson), each with a complete step-by-step worked example.
  Section 3 — Worked Examples: 2-4 fully solved exam-style problems with numbered steps,
               highlighted tricky parts, and common mistakes noted per example.
  Section 4 — Diagrams & Visual Descriptions: describe every diagram/figure in detail —
               what it shows, all labels, relationships between parts, and the key learning.
  Section 5 — Tips & Tricks for Exams: practical tips with mark allocation and frequency.
  Section 6 — Important Points to Remember: must-know facts and an exam checklist.

Write clear, simple English for Grade 6-13 students preparing for national exams.
"""

# ── DeepSeek API — two-pass approach ──────────────────────────────────────────
#
# Pass 1: identify the full lesson list from the guide (small response).
# Pass 2: generate full 6-section notes for ONE lesson at a time (avoids token
#          truncation that occurs when all lessons are requested in one shot).

LESSONS_LIST_TOOL = {
    "type": "function",
    "function": {
        "name": "list_lessons",
        "description": "List every lesson/chapter found in the teacher's guide.",
        "parameters": {
            "type": "object",
            "required": ["subject_title", "lessons"],
            "properties": {
                "subject_title": {"type": "string"},
                "lessons": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "required": ["lesson_number", "lesson_title"],
                        "properties": {
                            "lesson_number": {"type": "integer"},
                            "lesson_title":  {"type": "string"}
                        }
                    }
                }
            }
        }
    }
}


def _safe_json(raw):
    """Parse first valid JSON object from raw string, ignoring any trailing data."""
    obj, _ = json.JSONDecoder().raw_decode(raw.strip())
    return obj


def list_lessons_from_guide(guide_text, grade, subject_hint, client):
    """Pass 1 — return subject_title + list of {lesson_number, lesson_title}."""
    user_msg = (
        f"Grade: {grade}\n"
        f"Subject: {subject_hint or 'Unknown'}\n\n"
        "Read the teacher's guide and list every lesson/chapter using list_lessons.\n"
        "Use sequential lesson numbers (1, 2, 3...) in order of appearance.\n\n"
        "--- TEACHER'S GUIDE ---\n"
        f"{guide_text[:90000]}\n"
        "--- END ---"
    )
    resp = client.chat.completions.create(
        model="deepseek-chat",
        max_tokens=1024,
        messages=[
            {"role": "system", "content":
             "You are an expert at reading NIE Sri Lanka teacher's guides. "
             "List every lesson in order. Number them sequentially (1, 2, 3...). "
             "Always call the list_lessons function."},
            {"role": "user", "content": user_msg}
        ],
        tools=[LESSONS_LIST_TOOL],
        tool_choice={"type": "function", "function": {"name": "list_lessons"}}
    )
    msg = resp.choices[0].message
    if msg.tool_calls:
        return _safe_json(msg.tool_calls[0].function.arguments)
    raise ValueError("DeepSeek did not return the lesson list.")


def generate_one_lesson(guide_text, textbook_text, grade, subject_title,
                        lesson_number, lesson_title, client):
    """Pass 2 — generate full 6-section notes for a single lesson."""
    parts = [
        f"Grade: {grade}",
        f"Subject: {subject_title}",
        "",
        f"Generate detailed short notes for ONLY Lesson {lesson_number}: '{lesson_title}'.",
        f"Return exactly one entry in the lessons array with lesson_number={lesson_number}.",
        "",
        "--- TEACHER'S GUIDE (full — for context and content) ---",
        guide_text[:90000],
        "--- END TEACHER'S GUIDE ---",
    ]
    if textbook_text:
        parts += [
            "",
            "--- TEXTBOOK CONTENT (supplementary) ---",
            textbook_text[:30000],
            "--- END TEXTBOOK CONTENT ---",
        ]
    user_msg = "\n".join(parts)

    resp = client.chat.completions.create(
        model="deepseek-chat",
        max_tokens=8192,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": user_msg}
        ],
        tools=[NOTES_TOOL],
        tool_choice={"type": "function", "function": {"name": "create_short_notes"}}
    )
    msg = resp.choices[0].message
    if msg.tool_calls:
        data = _safe_json(msg.tool_calls[0].function.arguments)
        lessons = data.get("lessons", [])
        if lessons:
            return lessons[0]
        raise ValueError(f"No lesson data returned for lesson {lesson_number}.")
    raise ValueError("DeepSeek did not return the expected function call.")

# ── Word document helpers ──────────────────────────────────────────────────────

def _set_color(run, hex_color):
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )

def _heading(doc, text, level, color_hex=None, bold=True):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.bold = bold
        if color_hex:
            _set_color(run, color_hex)
    return p

def _body(doc, text, bold_prefix=False):
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix and ':' in text:
        idx = text.index(':') + 1
        run_b = p.add_run(text[:idx])
        run_b.bold = True
        p.add_run(text[idx:])
    else:
        p.add_run(text)

def _label_body(doc, label, content):
    """Bold label followed by normal content on same paragraph."""
    if not content:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(content)

def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(1)
    if ':' in text[:60]:
        idx = text.index(':') + 1
        r = p.add_run(text[:idx]); r.bold = True
        p.add_run(text[idx:])
    else:
        p.add_run(text)

def _numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text)

def _divider(doc, color='BDBDBD'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def _section_box_heading(doc, text, color_hex='283593'):
    """Section heading with light shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(12)
    _set_color(run, color_hex)
    # Light background shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'E8EAF6')
    pPr.append(shd)
    return p

# ── Build one lesson .docx ────────────────────────────────────────────────────

def build_lesson_doc(lesson, subject_title, grade, output_path):
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    ln = lesson.get('lesson_number', '?')
    lt = lesson.get('lesson_title', 'Lesson')

    # ── Header block ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(f"Grade {grade} — {subject_title}")
    r.bold = True; r.font.size = Pt(18)
    _set_color(r, '1A237E')

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Short Notes  |  NIE Sri Lanka Curriculum")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph()

    # Lesson title
    _heading(doc, f"Lesson {ln}: {lt}", level=1, color_hex='1A237E')
    _divider(doc, '1A237E')

    # ── Section 1: Core Concepts ──────────────────────────────────────────────
    concepts = lesson.get('core_concepts', [])
    if concepts:
        _section_box_heading(doc, "1. Core Concepts (Detailed Explanations)")
        for c in concepts:
            label = c.get('label', '')
            title = c.get('title', '')
            _heading(doc, f"{label}. {title}", level=3, color_hex='37474F')
            if c.get('definition'):
                _label_body(doc, "Definition", c['definition'])
            if c.get('explanation'):
                _label_body(doc, "Explanation", c['explanation'])
            if c.get('analogy'):
                _label_body(doc, "Think of it as", c['analogy'])
            if c.get('diagram_note'):
                _label_body(doc, "Diagram", c['diagram_note'])
            for conn in c.get('connections', []):
                _bullet(doc, f"Connected to: {conn}")

    # ── Section 2: Key Formulas / Definitions / Rules ─────────────────────────
    kf = lesson.get('key_formulas', {})
    revision_items = kf.get('revision', [])
    new_items      = kf.get('new', [])
    if revision_items or new_items:
        _section_box_heading(doc, "2. Key Formulas / Definitions / Rules")
        if revision_items:
            _heading(doc, "Revision (Previously Learned)", level=3, color_hex='546E7A')
            for item in revision_items:
                _body(doc, f"▶ {item.get('name','')}", bold_prefix=False)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.3)
                r = p.add_run("Formula/Rule: ")
                r.bold = True
                p.add_run(item.get('formula_or_rule', ''))
                if item.get('worked_example'):
                    _label_body(doc, "Example", item['worked_example'])
        if new_items:
            _heading(doc, "New (Introduced This Lesson)", level=3, color_hex='37474F')
            for item in new_items:
                _body(doc, f"▶ {item.get('name','')}", bold_prefix=False)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.3)
                r = p.add_run("Formula/Rule: ")
                r.bold = True
                p.add_run(item.get('formula_or_rule', ''))
                if item.get('worked_example'):
                    _label_body(doc, "Example", item['worked_example'])

    # ── Section 3: Worked Examples & Explanations ────────────────────────────
    examples = lesson.get('worked_examples', [])
    if examples:
        _section_box_heading(doc, "3. Worked Examples & Explanations")
        for ex in examples:
            pnum = ex.get('problem_number', '')
            prob = ex.get('problem', '')
            _heading(doc, f"Example {pnum}: {prob[:80]}{'…' if len(prob) > 80 else ''}", level=3, color_hex='37474F')
            if len(prob) > 80:
                _body(doc, prob)
            steps = ex.get('steps', [])
            if steps:
                p = doc.add_paragraph()
                r = p.add_run("Step-by-step solution:"); r.bold = True
                for step in steps:
                    _numbered(doc, step)
            for tp in ex.get('tricky_parts', []):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                r = p.add_run("⚠ Tricky: "); r.bold = True
                r.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
                p.add_run(tp)
            for cm in ex.get('common_mistakes', []):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                r = p.add_run("✗ Common mistake: "); r.bold = True
                r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
                p.add_run(cm)

    # ── Section 4: Diagrams & Visual Descriptions ─────────────────────────────
    diagrams = lesson.get('diagram_descriptions', [])
    if diagrams:
        _section_box_heading(doc, "4. Diagrams & Visual Descriptions")
        for d in diagrams:
            dnum = d.get('diagram_number', '')
            dtitle = d.get('title', '')
            _heading(doc, f"Diagram {dnum}: {dtitle}", level=3, color_hex='37474F')
            if d.get('what_it_shows'):
                _label_body(doc, "What it shows", d['what_it_shows'])
            labels = d.get('labels', [])
            if labels:
                p = doc.add_paragraph()
                r = p.add_run("Labels: "); r.bold = True
                for lbl in labels:
                    _bullet(doc, lbl, level=1)
            if d.get('relationships'):
                _label_body(doc, "Relationships", d['relationships'])
            if d.get('key_learning'):
                _label_body(doc, "Key learning", d['key_learning'])

    # ── Section 5: Tips & Tricks for Exams ───────────────────────────────────
    tips = lesson.get('exam_tips', [])
    if tips:
        _section_box_heading(doc, "5. Tips & Tricks for Exams")
        for tip in tips:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.add_run(tip.get('tip', ''))
            mark = tip.get('mark_info', '')
            freq = tip.get('frequency', '')
            if mark or freq:
                info_parts = []
                if mark: info_parts.append(mark)
                if freq: info_parts.append(freq)
                r = p.add_run(f"  [{' | '.join(info_parts)}]")
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                r.font.size = Pt(9)

    # ── Section 6: Important Points to Remember ───────────────────────────────
    imp = lesson.get('important_points', {})
    must_know  = imp.get('must_know', [])
    checklist  = imp.get('exam_checklist', [])
    if must_know or checklist:
        _section_box_heading(doc, "6. Important Points to Remember")
        if must_know:
            _heading(doc, "Must Know", level=3, color_hex='4A148C')
            for pt in must_know:
                _bullet(doc, pt)
        if checklist:
            _heading(doc, "Exam Checklist ✓", level=3, color_hex='1B5E20')
            for item in checklist:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Inches(0.3)
                r = p.add_run("☐  "); r.bold = True
                p.add_run(item)

    _divider(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"   📄 Saved: {output_path}")

# ── Main ───────────────────────────────────────────────────────────────────────

def safe_name(text):
    return re.sub(r'[<>:"/\\|?*]', '_', text).strip()

def extract_source(source, label, max_pages):
    """Extract and return combined text from a file or folder source."""
    if source is None:
        return ""
    pdfs = collect_pdfs(source)
    print(f"\n📖 Extracting {label} ({len(pdfs)} PDF(s))...")
    parts = []
    for pdf in pdfs:
        print(f"   • {pdf.name}")
        text = extract_pdf_text(pdf, max_pages=max_pages)
        if text:
            parts.append(f"=== {pdf.stem} ===\n{text}")
        else:
            print(f"   ⚠️  No text extracted (may be scanned image PDF)")
    combined = "\n\n".join(parts)
    if combined:
        print(f"   ✅ ~{len(combined.split()):,} words")
    return combined


def main():
    parser = argparse.ArgumentParser(
        description="NIE Sri Lanka — Short Notes Generator (6-Section Format)")

    # Teacher's guide source (primary — defines lesson structure)
    guide_group = parser.add_mutually_exclusive_group()
    guide_group.add_argument('--guide',        '-G',
                             help="Teacher's guide single PDF (primary lesson source)")
    guide_group.add_argument('--guide-folder', '-GF',
                             help="Folder of teacher's guide PDFs")

    # Textbook source (supplementary content)
    content_group = parser.add_mutually_exclusive_group()
    content_group.add_argument('--pdf',    '-f', help='Textbook single PDF (supplementary)')
    content_group.add_argument('--folder', '-F', help='Textbook folder (supplementary)')

    parser.add_argument('--grade',     '-g', required=True, help='Grade number (e.g. 10)')
    parser.add_argument('--medium',    '-m', default='english',
                        choices=['english', 'sinhala', 'tamil'],
                        help='Medium: english / sinhala / tamil (default: english)')
    parser.add_argument('--subject',   '-s', default=None,
                        help='Subject name hint (auto-detected if omitted)')
    parser.add_argument('--output',    '-o',
                        default=r'C:\Users\shaki\Downloads\nie-downloader\output\short-notes',
                        help='Base output root (default: output/short-notes/)')
    parser.add_argument('--api-key',   '-k', default=None,
                        help='DeepSeek API key (or set DEEPSEEK_API_KEY env variable)')
    parser.add_argument('--max-pages', type=int, default=None,
                        help='Limit pages extracted per PDF')
    parser.add_argument('--lesson', type=int, default=None,
                        help='Retry/generate notes for a single lesson number only')
    parser.add_argument('--lesson-title', default=None,
                        help='Title for --lesson (required when using --lesson)')
    args = parser.parse_args()

    # Validate: at least one source must be provided
    guide_source   = args.guide or args.guide_folder
    content_source = args.pdf   or args.folder
    if not guide_source and not content_source:
        parser.error("Provide at least --guide / --guide-folder (and optionally --pdf / --folder).")

    api_key = args.api_key or os.environ.get('DEEPSEEK_API_KEY')
    if not api_key:
        print("❌ No API key. Use --api-key or set DEEPSEEK_API_KEY.")
        sys.exit(1)

    print("=" * 60)
    print("📝 NIE Short Notes Generator  (6-Section Expanded Format)")
    print("=" * 60)
    print(f"   Grade:       {args.grade}")
    print(f"   Medium:      {args.medium}")
    if guide_source:
        print(f"   Guide:       {guide_source}")
    if content_source:
        print(f"   Textbook:    {content_source}")
    print(f"   Output root: {Path(args.output).resolve()}")
    print("=" * 60)

    # Extract guide text (primary)
    guide_text = ""
    if guide_source:
        guide_text = extract_source(guide_source, "teacher's guide", args.max_pages)
        if not guide_text:
            print("❌ No text extracted from teacher's guide. Scanned PDFs require OCR.")
            sys.exit(1)

    # Extract textbook text (supplementary, optional)
    textbook_text = ""
    if content_source:
        textbook_text = extract_source(content_source, "textbook", args.max_pages)

    # Subject hint
    subject_hint = args.subject
    if not subject_hint:
        primary = guide_source or content_source
        subject_hint = Path(primary).stem if (args.guide or args.pdf) else Path(primary).name

    # Create shared DeepSeek client
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

    # ── Single-lesson retry mode ───────────────────────────────────────────────
    if args.lesson is not None:
        if not args.lesson_title:
            parser.error("--lesson-title is required when using --lesson")
        ln = args.lesson
        lt = args.lesson_title
        subject_title = subject_hint or 'Subject'
        subject_dir   = (Path(args.output)
                         / f"{args.medium}-medium"
                         / f"grade {args.grade}"
                         / safe_name(subject_title))
        subject_dir.mkdir(parents=True, exist_ok=True)
        print(f"\n🔁 Retrying Lesson {ln:02d}: {lt}")
        print(f"   🤖 Sending to DeepSeek...", end='', flush=True)
        try:
            lesson_data = generate_one_lesson(
                guide_text, textbook_text, args.grade, subject_title, ln, lt, client)
            print(" done")
            filename = f"Lesson {ln:02d} - {safe_name(lt)}.docx"
            build_lesson_doc(lesson_data, subject_title, args.grade, subject_dir / filename)
        except Exception as e:
            print(f" ❌ Failed: {e}")
            sys.exit(1)
        print(f"\n{'=' * 60}")
        print(f"✅ Done!  Lesson {ln:02d} saved  →  {subject_dir}/")
        print(f"{'=' * 60}")
        return

    # ── Pass 1: identify lesson list from the teacher's guide ─────────────────
    print(f"\n🔍 Pass 1 — identifying lessons from teacher's guide...")
    try:
        lesson_list_data = list_lessons_from_guide(
            guide_text, args.grade, subject_hint, client)
    except Exception as e:
        print(f"\n❌ DeepSeek API error (lesson list): {e}")
        sys.exit(1)

    subject_title = lesson_list_data.get('subject_title', subject_hint or 'Subject')
    lesson_list   = lesson_list_data.get('lessons', [])
    # Normalise lesson_number to int; renumber sequentially if all identical
    for l in lesson_list:
        l['lesson_number'] = int(l['lesson_number'])
    nums = [l['lesson_number'] for l in lesson_list]
    if len(set(nums)) == 1:
        for i, l in enumerate(lesson_list, 1):
            l['lesson_number'] = i
    print(f"   ✅ Found {len(lesson_list)} lesson(s) in '{subject_title}'")
    for l in lesson_list:
        print(f"      Lesson {l['lesson_number']:02d}: {l['lesson_title']}")

    if not lesson_list:
        print("❌ No lessons identified. Check the teacher's guide PDF.")
        sys.exit(1)

    # Set up output folder: <root>/<medium>-medium/grade <N>/<subject>/
    subject_dir = (Path(args.output)
                   / f"{args.medium}-medium"
                   / f"grade {args.grade}"
                   / safe_name(subject_title))
    subject_dir.mkdir(parents=True, exist_ok=True)

    # ── Pass 2: generate notes one lesson at a time ───────────────────────────
    print(f"\n📝 Pass 2 — generating notes lesson by lesson...")
    all_lessons_json = []
    success = 0
    failed  = 0

    for i, l in enumerate(lesson_list, 1):
        ln = l['lesson_number']
        lt = l['lesson_title']
        print(f"\n[{i}/{len(lesson_list)}] Lesson {ln:02d}: {lt}")
        print(f"   🤖 Sending to DeepSeek...", end='', flush=True)
        try:
            lesson_data = generate_one_lesson(
                guide_text, textbook_text,
                args.grade, subject_title,
                ln, lt, client)
            print(" done")
            all_lessons_json.append(lesson_data)
            filename = f"Lesson {ln:02d} - {safe_name(lt)}.docx"
            out_path = subject_dir / filename
            build_lesson_doc(lesson_data, subject_title, args.grade, out_path)
            success += 1
        except Exception as e:
            print(f" ❌ Failed: {e}")
            failed += 1

    # Save combined JSON alongside the lesson docs
    notes_data = {"subject_title": subject_title, "lessons": all_lessons_json}
    json_file  = subject_dir / f"Grade {args.grade} - {safe_name(subject_title)} - notes.json"
    with open(json_file, 'w', encoding='utf-8') as jf:
        json.dump(notes_data, jf, ensure_ascii=False, indent=2)
    print(f"\n   🗂️  JSON: {json_file}")

    print(f"\n{'=' * 60}")
    print(f"✅ Done!  {success} saved, {failed} failed  →  {subject_dir}/")
    print(f"{'=' * 60}")

if __name__ == '__main__':
    main()
